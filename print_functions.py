# Print Abstract sections to document

import word_features
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn


def print_person(person):
    institute_names = []
    for institute in person['institutes']:
        institute_names.append(institute['name'])
    institute_names = " (" + (", ".join(institute_names)) + ")"
    return person['initials'] + " " + person['last_name'] + institute_names


def print_program(d, sessions):
    word_features.add_page_numbers("Contents", d.sections[-1])
    word_features.add_alternating_headers("Program", "", d.sections[-1])

    d.add_heading('Program', 0)

    for session in sessions:
        # Date + Time
        p = d.add_paragraph()
        p.add_run(session['date'] + " " + session['start'].strftime("%I:%M") + " - " + session['end'].strftime("%I:%M")).bold = True
        p.paragraph_format.keep_with_next = True
        # Time
        p = d.add_paragraph()
        p.paragraph_format.keep_with_next = True
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(3.5))
        p.paragraph_format.first_line_indent = Cm(-3.5)
        p.paragraph_format.left_indent = Cm(3.5)

        p.add_run(session['start'].strftime("%I:%M") + " - " + session['end'].strftime("%I:%M")).bold = True
        r = p.add_run()
        r.add_tab()
        r.add_text(session['code'] + " - " + session['name'])

        if session['chair'] is not None:
            r.add_break()
            r.add_text("Session Chair: " + print_person(session['chair']))

        for paper in session['papers']:
            p = d.add_paragraph()
            p.paragraph_format.keep_with_next = True
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(3.5))
            p.paragraph_format.first_line_indent = Cm(-3.5)
            p.paragraph_format.left_indent = Cm(3.5)

            p.add_run(paper['code']).bold = True
            r = p.add_run()

            r.add_tab()
            r.add_text(paper['title'])
            r.add_break()
            for author in paper['authors']:
                r.add_text(print_person(author))
        p.paragraph_format.keep_with_next = False


def print_abstracts(d, sessions):
    for session in sessions:
        session_section = d.add_section(WD_SECTION.NEW_PAGE)
        # Header
        name = session['code'] + " - " + session['name']
        time = session['date'] + " " + session['start'].strftime("%I:%M") + " - " + session['end'].strftime("%I:%M")
        session_section.header.is_linked_to_previous = False
        session_section.even_page_header.is_linked_to_previous = False
        word_features.add_alternating_headers(session['name'], time, session_section)
        d.add_heading(name, 0)

        # Footer
        chair = ""
        session_section.footer.is_linked_to_previous = False
        session_section.even_page_footer.is_linked_to_previous = False
        if session['chair'] is not None:
            chair = "Chair: " + print_person(session['chair'])
        word_features.add_page_numbers(chair, session_section)

        for paper in session['papers']:
            p = d.add_paragraph()
            title_run = p.add_run(paper['title'])
            title_run.add_break()
            title_run.bold = True
            authors_run = p.add_run()
            for author in paper['authors']:
                authors_run.add_text(print_person(author))
            authors_run.add_break()
            authors_run.italic = True

            p.add_run(paper['abstract'])
            if paper['funding'] is not None:
                funding_run = p.add_run()
                funding_run.add_break()
                funding_run.add_text("Funding: " + paper['funding'])
                funding_run.bold = True

            if paper['footnote'] is not None:
                footnote_run = p.add_run()
                footnote_run.add_break()
                footnote_run.add_text(paper['footnote'])


def is_primary_author(paper, person):
    all_authors = paper['all_authors']
    for author in all_authors:
        if author['id'] == person['id'] and \
                'person_type' in author and \
                author['person_type'] == "Primary Author":
            return True
    return False

def print_authors(d, authors):
    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    author_section = d.add_section(WD_SECTION.NEW_PAGE)

    d.add_heading("Author Index", 0)
    # Header
    author_section.header.is_linked_to_previous = False
    author_section.even_page_header.is_linked_to_previous = False
    word_features.add_alternating_headers("Author Index", "", author_section)

    author_section.footer.is_linked_to_previous = False
    author_section.even_page_footer.is_linked_to_previous = False
    word_features.add_page_numbers("Italic papercodes indicate primary author", author_section)

    sectPr = author_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    for letter in alphabet:

        author_subset = list(filter(lambda author: author['last_name'][0] == letter, authors))

        if len(author_subset) > 0:
            d.add_heading(letter, 1)
            for author in author_subset:
                p = d.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Cm(3.5))
                p.paragraph_format.first_line_indent = Cm(-3.5)
                p.paragraph_format.left_indent = Cm(3.5)
                author_run = p.add_run(author['last_name'] + ", " + author['initials'])
                author_run.add_tab()
                for i, paper in enumerate(author['papers']):
                    joiner = ""
                    if i != len(author['papers']) - 1:
                        joiner = ", "
                    p.add_run(paper['code'] + joiner).italic = is_primary_author(paper, author)

