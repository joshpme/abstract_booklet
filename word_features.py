# SPECIAL WORD FEATURES (Like page numbers)

from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, ns


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)
    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)
    run._r.append(fldEnd)


def add_page_numbers(text, section):
    odd_para = section.footer.add_paragraph()
    odd_run = odd_para.add_run()
    even_para = section.even_page_footer.add_paragraph()
    even_run = even_para.add_run()

    odd_run.add_text(text)
    odd_run.add_tab()
    add_page_number(odd_run)
    odd_run.bold = True
    add_page_number(even_run)
    even_run.add_tab()
    even_run.add_text(text)
    even_run.bold = True

    odd_stops = odd_para.paragraph_format.tab_stops
    even_stops = even_para.paragraph_format.tab_stops
    odd_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.RIGHT)
    even_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.RIGHT)


def add_alternating_headers(left, right, section):
    odd_para = section.header.add_paragraph()
    odd_run = odd_para.add_run()
    odd_run.add_text(left)
    odd_run.add_tab()
    odd_run.add_text(right)
    odd_run.bold = True

    even_para = section.even_page_header.add_paragraph()
    even_run = even_para.add_run()
    even_run.add_text(right)
    even_run.add_tab()
    even_run.add_text(left)
    even_run.bold = True

    odd_stops = odd_para.paragraph_format.tab_stops
    even_stops = even_para.paragraph_format.tab_stops
    odd_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.RIGHT)
    even_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.RIGHT)
